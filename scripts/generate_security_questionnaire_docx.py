from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from security_questionnaire_source import (
    DOCUMENT_METADATA,
    DOCKER_ANNEX_ROWS,
    EVIDENCE_INDEX,
    QUESTIONNAIRE_ROWS,
    SYNC_ANNEX_ROWS,
    VALIDATION_EXPECTATIONS,
    VISUAL_SELECTION,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "doc"
EVIDENCE_DIR = OUTPUT_DIR / "evidence"
DOCX_PATH = OUTPUT_DIR / "respuesta_cuestionario_seguridad_onprem_sura.docx"
SOURCE_JSON_PATH = OUTPUT_DIR / "fuente_cuestionario_seguridad_onprem.json"
EVIDENCE_INDEX_PATH = EVIDENCE_DIR / "index.json"


def ensure_output_dirs() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)


def set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CBD5E1")
        borders.append(border)
    tbl_pr.append(borders)


def style_paragraph(paragraph, *, bold: bool = False, size: float = 9, color: str = "0F172A", align=None) -> None:
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    size: float = 9,
    color: str = "0F172A",
    align=None,
):
    paragraph = doc.add_paragraph()
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(4)
    style_paragraph(paragraph, bold=bold, size=size, color=color, align=align)
    return paragraph


def add_bullet(doc: Document, text: str, *, size: float = 8.8, color: str = "0F172A") -> None:
    paragraph = doc.add_paragraph(style=None)
    paragraph.style = doc.styles["List Bullet"]
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(2)
    style_paragraph(paragraph, size=size, color=color)


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8") if path.exists() else ""


def parse_validation_results() -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for item in VALIDATION_EXPECTATIONS:
        path = ROOT / item["evidence"]
        content = read_text(path)
        ok = path.exists() and all(marker in content for marker in item["success_markers"])
        summary = item["summary"]

        if item["command"] == "pnpm lint":
            match = re.search(r"(\d+)\s+errors,\s+(\d+)\s+warnings", content)
            if match:
                summary = f"{summary} Resultado observado: {match.group(1)} errores y {match.group(2)} advertencias."
        elif item["command"] == "pnpm test:e2e":
            passed = re.search(r"# pass\s+(\d+)", content)
            failed = re.search(r"# fail\s+(\d+)", content)
            if passed and failed:
                summary = f"{summary} Resultado observado: {passed.group(1)} pruebas aprobadas y {failed.group(1)} fallas."

        rows.append(
            {
                "command": item["command"],
                "status": "Exitoso" if ok else "Revisar evidencia",
                "evidence": item["evidence"],
                "summary": summary,
            }
        )
    return rows


def export_json_artifacts() -> None:
    payload = {
        "metadata": DOCUMENT_METADATA,
        "rows": QUESTIONNAIRE_ROWS,
        "evidence_index": EVIDENCE_INDEX,
        "docker_annex": DOCKER_ANNEX_ROWS,
        "sync_annex": SYNC_ANNEX_ROWS,
        "visual_selection": VISUAL_SELECTION,
    }
    SOURCE_JSON_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    EVIDENCE_INDEX_PATH.write_text(json.dumps(EVIDENCE_INDEX, ensure_ascii=False, indent=2), encoding="utf-8")


def add_header(cell, text: str) -> None:
    cell.text = text
    set_cell_shading(cell, "1D4ED8")
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        if not paragraph.runs:
            paragraph.add_run(text)
        style_paragraph(paragraph, bold=True, size=8.2, color="FFFFFF")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def fill_cell(cell, text: str, *, size: float = 8, bold: bool = False, color: str = "111827", align=None) -> None:
    cell.text = ""
    blocks = text.split("\n")
    for index, block in enumerate(blocks):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.text = block
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        style_paragraph(paragraph, bold=bold, size=size, color=color, align=align)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def make_table(doc: Document, headers: list[str], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for cell, header, width in zip(header_row.cells, headers, widths):
        add_header(cell, header)
        cell.width = Inches(width)
    return table


def build_questionnaire_evidence_text(row: dict[str, object]) -> str:
    summary = row.get("evidencia_resumen")
    if isinstance(summary, str) and summary.strip():
        return summary

    evidence_ids = row["evidence_ids"][:5]
    visuals = [item["id"] for item in row["evidencia_visual"][:2]]
    merged = []
    for evidence_id in evidence_ids:
        merged.append(evidence_id)
    for evidence_id in visuals:
        if evidence_id not in merged:
            merged.append(evidence_id)
    return "\n".join(merged) if merged else "Pendiente de anexo del contratista"


def append_cover(doc: Document) -> None:
    counts = Counter(row["respuesta"] for row in QUESTIONNAIRE_ROWS)

    add_paragraph(doc, DOCUMENT_METADATA["title"], bold=True, size=20, color="1E3A8A", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, DOCUMENT_METADATA["subtitle"], bold=True, size=11, color="475569", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"Destinatario: {DOCUMENT_METADATA['client']}", size=10, color="0F172A", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"Fecha de emision: {DOCUMENT_METADATA['date']}", size=10, color="0F172A", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"Repositorio analizado: {DOCUMENT_METADATA['repository']}", size=8.5, color="475569", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_paragraph(doc, "Declaracion ejecutiva", bold=True, size=12, color="1E3A8A")
    add_paragraph(doc, DOCUMENT_METADATA["scope_note"], size=9, color="334155")
    for bullet in DOCUMENT_METADATA["executive_summary"]:
        add_bullet(doc, bullet)

    add_paragraph(doc, "Resumen de respuestas", bold=True, size=12, color="1E3A8A")
    table = make_table(doc, ["Sí", "Se adjuntará posteriormente", "Total"], [1.8, 2.8, 1.6])
    cells = table.add_row().cells
    fill_cell(cells[0], str(counts.get("Sí", 0)), size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(cells[1], str(counts.get("Se adjuntará posteriormente", 0)), size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(cells[2], str(len(QUESTIONNAIRE_ROWS)), size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def append_questionnaire_table(doc: Document) -> None:
    doc.add_page_break()
    add_paragraph(doc, "Tabla Integral del Cuestionario", bold=True, size=12, color="1E3A8A")
    add_paragraph(
        doc,
        "La matriz conserva la numeración original 1-54. Los controles técnicos se responden con evidencia implementada y los controles organizacionales o certificatorios quedan marcados para anexo posterior del contratista.",
        size=9,
        color="334155",
    )
    table = make_table(
        doc,
        ["N°", "Sección", "Tipo", "Pregunta/Control", "Respuesta", "Comentarios", "Evidencia"],
        [0.42, 1.15, 0.72, 3.05, 0.95, 3.15, 1.25],
    )

    for item in QUESTIONNAIRE_ROWS:
        cells = table.add_row().cells
        control_text = f"{item['control']}\n{item['pregunta']}"
        evidence_text = build_questionnaire_evidence_text(item)
        fill_cell(cells[0], str(item["numero_control"]), size=8.4, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(cells[1], item["seccion"], size=7.7)
        fill_cell(cells[2], item["tipo_control"], size=7.7, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(cells[3], control_text, size=7.5)
        fill_cell(cells[4], item["respuesta"], size=8.1, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(cells[5], item["comentario_ejecutivo"], size=7.4)
        fill_cell(cells[6], evidence_text, size=7.4)


def append_evidence_annex(doc: Document) -> None:
    doc.add_page_break()
    add_paragraph(doc, "Anexo A. Matriz de Evidencia", bold=True, size=12, color="1E3A8A")
    add_paragraph(doc, "Cada evidencia se indexa con identificador EV, tipo, archivo y controles relacionados.", size=9, color="334155")
    table = make_table(doc, ["ID", "Tipo", "Descripcion", "Archivo", "Controles"], [0.55, 0.85, 3.60, 3.75, 1.15])
    for item in EVIDENCE_INDEX:
        cells = table.add_row().cells
        controls_text = ", ".join(str(control) for control in item["controles_relacionados"])
        fill_cell(cells[0], item["id"], size=8.0, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(cells[1], item["tipo"], size=7.8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(cells[2], item["descripcion"], size=7.6)
        fill_cell(cells[3], item["archivo"], size=7.2)
        fill_cell(cells[4], controls_text, size=7.3)


def append_docker_annex(doc: Document) -> None:
    doc.add_page_break()
    add_paragraph(doc, "Anexo B. Arquitectura Docker y Operacion On-Premise", bold=True, size=12, color="1E3A8A")
    add_paragraph(
        doc,
        "La implantacion deja una topologia completa con aplicacion, PostgreSQL, proxy TLS y backup-runner dentro del perimetro del cliente.",
        size=9,
        color="334155",
    )
    table = make_table(doc, ["Componente", "Implementacion", "Evidencia"], [1.60, 5.25, 3.25])
    for item in DOCKER_ANNEX_ROWS:
        cells = table.add_row().cells
        evidence_text = "\n".join(item["evidence_ids"])
        fill_cell(cells[0], item["component"], size=8.1, bold=True)
        fill_cell(cells[1], item["implementation"], size=7.8)
        fill_cell(cells[2], evidence_text, size=7.8)

    add_paragraph(doc, "Controles destacados", bold=True, size=10, color="1E3A8A")
    highlights = [
        "Despliegue completamente on-premise con componentes desacoplados y healthchecks por servicio.",
        "TLS interno gestionado por proxy dedicado, compatible con certificados corporativos.",
        "Persistencia central en PostgreSQL y continuidad local mediante cola de sincronizacion persistente.",
        "Respaldo y restauracion verificados con checksum, manifiesto y base restaurada.",
    ]
    for item in highlights:
        add_bullet(doc, item)


def append_sync_annex(doc: Document) -> None:
    doc.add_page_break()
    add_paragraph(doc, "Anexo C. Continuidad Local-First y Sincronizacion", bold=True, size=12, color="1E3A8A")
    add_paragraph(
        doc,
        "Esta capa fue implementada para demostrar continuidad operativa sin conectividad y sincronizacion posterior contra un repositorio central on-premise.",
        size=9,
        color="334155",
    )
    table = make_table(doc, ["Capacidad", "Implementacion", "Evidencia"], [1.70, 5.20, 3.20])
    for item in SYNC_ANNEX_ROWS:
        cells = table.add_row().cells
        fill_cell(cells[0], item["component"], size=8.0, bold=True)
        fill_cell(cells[1], item["implementation"], size=7.7)
        fill_cell(cells[2], "\n".join(item["evidence_ids"]), size=7.7)

    add_paragraph(doc, "Puntos de control del flujo hibrido", bold=True, size=10, color="1E3A8A")
    checkpoints = [
        "Bootstrap inicial con sesion server-side y registro de dispositivo.",
        "Creacion y modificacion de registros en local con cola persistente.",
        "Pull y push incremental hacia PostgreSQL on-premise.",
        "Deteccion de conflicto por version base y preservacion del caso para resolucion explicita.",
        "Exportacion de eventos de seguridad para monitoreo y correlacion central.",
    ]
    for item in checkpoints:
        add_bullet(doc, item)


def append_validation_annex(doc: Document, validation_rows: list[dict[str, str]]) -> None:
    doc.add_page_break()
    add_paragraph(doc, "Anexo D. Validaciones Ejecutadas", bold=True, size=12, color="1E3A8A")
    add_paragraph(doc, "Los siguientes resultados fueron obtenidos y conservados como evidencia reproducible.", size=9, color="334155")
    table = make_table(doc, ["Comando", "Estado", "Archivo de evidencia", "Resumen"], [2.25, 1.00, 3.10, 4.25])
    for item in validation_rows:
        cells = table.add_row().cells
        fill_cell(cells[0], item["command"], size=8.0, bold=True)
        fill_cell(cells[1], item["status"], size=8.0, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(cells[2], item["evidence"], size=7.4)
        fill_cell(cells[3], item["summary"], size=7.5)


def append_operational_annex(doc: Document) -> None:
    doc.add_page_break()
    add_paragraph(doc, "Anexo E. Controles con Llenado Posterior", bold=True, size=12, color="1E3A8A")
    add_paragraph(
        doc,
        "Los siguientes controles dependen de evidencia organizacional, certificatoria, física o contractual que no forma parte de este paquete técnico y deberá adjuntarse posteriormente.",
        size=9,
        color="334155",
    )
    pending_rows = [row for row in QUESTIONNAIRE_ROWS if row["respuesta"] == "Se adjuntará posteriormente"]
    table = make_table(doc, ["N°", "Sección", "Pregunta/Control", "Estado"], [0.55, 1.55, 6.45, 2.00])
    for row in pending_rows:
        cells = table.add_row().cells
        fill_cell(cells[0], str(row["numero_control"]), size=8.0, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(cells[1], row["seccion"], size=7.6)
        fill_cell(cells[2], f"{row['control']}\n{row['pregunta']}", size=7.4)
        fill_cell(cells[3], "Se adjuntará posteriormente", size=7.6, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    notes = [
        "Este anexo evita afirmar documentación no adjunta y permite completar la respuesta final del contratista con precisión.",
        "La modalidad on-premise híbrida mantiene el control del entorno dentro del cliente y reduce dependencia de terceros no gestionados.",
        "Las pruebas y anexos corporativos podrán agregarse posteriormente sin alterar la matriz técnica ya validada.",
    ]
    for note in notes:
        add_bullet(doc, note)


def append_visual_annex(doc: Document) -> None:
    doc.add_page_break()
    add_paragraph(doc, "Anexo F. Evidencia Visual Seleccionada", bold=True, size=12, color="1E3A8A")
    add_paragraph(doc, "Se incluyen screenshots funcionales y operativos para reforzar la confianza sobre la solucion implementada.", size=9, color="334155")

    for evidence_id in VISUAL_SELECTION:
        item = next(entry for entry in EVIDENCE_INDEX if entry["id"] == evidence_id)
        image_path = ROOT / item["archivo"]
        add_paragraph(doc, f"{item['id']} - {item['descripcion']}", bold=True, size=9.5, color="1E293B")
        if image_path.exists():
            doc.add_picture(str(image_path), width=Inches(8.9))
        else:
            add_paragraph(doc, f"Imagen no disponible al momento de la generacion: {item['archivo']}", size=8.2, color="B91C1C")
        add_paragraph(doc, f"Archivo: {item['archivo']}", size=8, color="475569")


def build_document() -> None:
    ensure_output_dirs()
    export_json_artifacts()

    document = Document()
    set_landscape(document.sections[0])
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal_style.font.size = Pt(9)

    validation_rows = parse_validation_results()

    append_cover(document)
    append_questionnaire_table(document)
    append_evidence_annex(document)
    append_docker_annex(document)
    append_sync_annex(document)
    append_validation_annex(document, validation_rows)
    append_operational_annex(document)
    append_visual_annex(document)

    document.save(DOCX_PATH)


if __name__ == "__main__":
    build_document()
